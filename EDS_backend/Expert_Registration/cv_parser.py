from docx import Document
import re
import pprint
from dateutil import parser
from dateutil.parser import parse as date_parse
import datetime # Import datetime for date objects

def safe_date_parse(date_str):
    if not date_str:
        return None

    original = date_str
    date_str = date_str.strip()
    if not date_str:
        return None

    # Normalize
    date_str = date_str.replace("–", "-").replace("—", "-")
    date_str = date_str.replace("/", "-")
    date_str = re.sub(r"\b(from|to|and|through)\b", "", date_str, flags=re.IGNORECASE)
    date_str = re.sub(r"\bSept\b", "September", date_str)
    date_str = re.sub(r"\b2O(\d{2})\b", r"20\1", date_str)  # Fix 2O21 → 2021
    date_str = re.sub(r"(\d+)\s*-\s*(\d+)", r"\1 to \2", date_str)  # 7-29 → 7 to 29
    date_str = re.sub(r"[^\w\s,-]", "", date_str)  # Strip punctuation

    # Fragmented month names like "Oc", "ber"
    months = ['January','February','March','April','May','June','July','August','September','October','November','December']
    for month in months:
        if month[:3].lower() in date_str.lower():
            date_str = re.sub(r"\b" + re.escape(month[:3]) + r"\b", month, date_str, flags=re.IGNORECASE)

    # Special case: "March 7 to 29 2014"
    match = re.match(r"^(?P<month>\w+)\s+(?P<day1>\d{1,2})\s+to\s+(?P<day2>\d{1,2})\s+(?P<year>\d{4})$", date_str)
    if match:
        try:
            return datetime.date(
                int(match.group("year")),
                datetime.datetime.strptime(match.group("month"), "%B").month,
                int(match.group("day1"))
            )
        except Exception:
            pass

    # "January to June 2000"
    match = re.match(r"^(?P<month1>\w+)\s+to\s+(?P<month2>\w+)\s+(?P<year>\d{4})$", date_str)
    if match:
        try:
            return datetime.date(
                int(match.group("year")),
                datetime.datetime.strptime(match.group("month1"), "%B").month,
                1
            )
        except Exception:
            pass

    # "June 2000"
    match = re.match(r"^(?P<month>\w+)\s+(?P<year>\d{4})$", date_str)
    if match:
        try:
            return datetime.date(
                int(match.group("year")),
                datetime.datetime.strptime(match.group("month"), "%B").month,
                1
            )
        except Exception:
            pass

    # Handle day-first fallback: "31, 2010" or "15. 2006"
    match = re.match(r"(?P<day>\d{1,2})[.,]?\s*(?P<year>\d{4})", date_str)
    if match:
        try:
            return datetime.date(
                int(match.group("year")),
                1,
                int(match.group("day"))
            )
        except Exception:
            pass

    # Handle 'present'
    if "present" in date_str.lower() or "continues up to" in date_str.lower():
        return datetime.date.today()

    # Fuzzy general parse
    try:
        return date_parse(date_str, fuzzy=True, default=datetime.datetime(1900, 1, 1)).date()
    except Exception as e:
        # Last ditch: just extract a year if we can
        year_match = re.search(r'\b(19|20)\d{2}\b', date_str)
        if year_match:
            return datetime.date(int(year_match.group()), 1, 1)

        print(f"[!] Failed parsing: {repr(original)} — Error: {str(e)}")
        return None

def parse_date_range(period_str: str):
    import re

    # Normalize punctuation and spaces
    period_str = period_str.replace(",", " ").replace("–", "-").replace("—", "-").strip()
    period_str = re.sub(r"\s*-\s*", "-", period_str)
    period_str = re.sub(r"\s+", " ", period_str)

    # Month first: March 12-31 2012
    match = re.match(r"^([A-Za-z]+)\s+(\d{1,2})-(\d{1,2})\s+(\d{4})$", period_str)
    if match:
        month, day1, day2, year = match.groups()
        start = safe_date_parse(f"{month} {day1}, {year}")
        end = safe_date_parse(f"{month} {day2}, {year}")
        return start, end

    # Day first with dash: 22-29 May 2009
    match = re.match(r"^(\d{1,2})-(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})$", period_str)
    if match:
        day1, day2, month, year = match.groups()
        start = safe_date_parse(f"{month} {day1}, {year}")
        end = safe_date_parse(f"{month} {day2}, {year}")
        return start, end

    # Day first with "to": 4 to 15 August 2008
    match = re.match(r"^(\d{1,2})\s+to\s+(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})$", period_str, flags=re.IGNORECASE)
    if match:
        day1, day2, month, year = match.groups()
        start = safe_date_parse(f"{month} {day1}, {year}")
        end = safe_date_parse(f"{month} {day2}, {year}")
        return start, end

    # Your existing split logic here
    parts = re.split(r"\s+(?:to|and|continues up to)\s+|\s*-\s*|\/", period_str, flags=re.IGNORECASE)
    if len(parts) == 2:
        raw_start, raw_end = parts[0].strip(), parts[1].strip()
        end = safe_date_parse(raw_end)
        start = safe_date_parse(raw_start)
        if start and start.year == 1900 and end and end.year != 1900:
            try:
                start = safe_date_parse(f"{raw_start} {end.year}")
            except:
                pass
        return start, end

    single = safe_date_parse(period_str)
    return single, single




def normalize_level(level_raw):
    """Normalizes education level strings to a consistent format."""
    level_raw = level_raw.lower()
    if "phd" in level_raw:
        return "phd"
    elif "master" in level_raw:
        return "masters"
    elif "diploma" in level_raw:
        return "diploma"
    elif "degree" in level_raw:
        return "degree"
    return "degree" # Default for unparsed levels

def extract_personal_info(full_text, doc):
    """
    Extracts personal and contact information, prioritizing table searches
    for "Position Title", "Name", "Date of Birth", "Country of Citizenship/Residence".
    Falls back to regex on full text for contact info if not found in tables.
    """
    personal_data = {
        "first_name": None,
        "last_name": None,
        "email": None,
        "phone_number": None,
        "cv_language": "English",
        "country": None,
        "date_of_birth": None,
        "gender": None, # Gender not present in sample CV
        "current_position": None,
        "name_suffix": None # Name suffix not present in sample CV
    }

    # Flags to track if information is found in tables
    found_in_table = {
        "name": False,
        "dob": False,
        "country": False,
        "position": False
    }

    # --- Step 1: Prioritize extraction from tables for initial personal details ---
    # Iterate through all tables to find the personal details table
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
          
            # Heuristic to identify the initial personal details table (2 columns, specific keywords)
            if len(cells) == 2:
                key_cell = cells[0].strip().lower()
                value_cell = cells[1].strip()

                if "position title" in key_cell and not found_in_table["position"]:
                    personal_data["current_position"] = value_cell
                    found_in_table["position"] = True
                elif "name of expert" in key_cell and not found_in_table["name"]:
                        # Extract suffix inside parentheses (e.g., PhD, MSc)
                    suffix_match = re.search(r"\((.*?)\)", value_cell)
                    suffix = suffix_match.group(1).strip() if suffix_match else None
                    raw_name = re.sub(r"\(.*?\)", "", value_cell).replace("Dr.", "").strip()
                    parts = raw_name.split()
              
                    if len(parts) >= 2:

                        personal_data["first_name"] = parts[0]
                        personal_data["last_name"] = " ".join(parts[1:])
                        personal_data['name_suffix']=suffix

                    elif parts:
                        personal_data["first_name"] = parts[0]
                    found_in_table["name"] = True
                elif "date of birth" in key_cell and not found_in_table["dob"]:
                    
                    personal_data["date_of_birth"] = safe_date_parse(value_cell.replace(',', ''))
                    found_in_table["dob"] = True
                elif "country of citizenship/residence" in key_cell and not found_in_table["country"]:
                    personal_data["country"] = value_cell
                    found_in_table["country"] = True
            
            # If all key fields are found, no need to process further tables for these
            if all(found_in_table.values()):
                break
        if all(found_in_table.values()):
            break

    # --- Step 2: Extract Email and Phone, typically found in plain text at the end ---
    # These are explicitly mentioned as "just text format" at the end of the document
    # (after "Expert's contact information")
    
    # Try to find email from the full text
    email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]+", full_text)
    if email_match:
        personal_data["email"] = email_match.group()

    # Try to find phone from the full text
    phone_match = re.search(r"(?:Cell\s*phone|Phone)[:：]?\s*(\+?\d[\d\s-]{8,}\d)", full_text, re.IGNORECASE)
    if phone_match:
        personal_data["phone_number"] = re.sub(r"\s+|-", "", phone_match.group(1))

    return personal_data


def parse_publications(full_text):
    """
    Extracts publications, journals, books, and learning modules from text blocks.
    Processes sections sequentially to avoid content bleeding.
    """
    publications_data = {
        "journals": [],
        "books": [],
        "learning_module": [],
        "publications": [] # Combined unique publications
    }

    # Define robust markers for different publication sections
    # Using non-greedy patterns and clear boundaries.
    # The regex will capture everything from the start marker until the next *known* section marker or end of text.
    
    # Initial section markers (should be distinct)
    journals_start_pattern = r"Journals:"
    book_chapters_start_pattern = r"Book chapters, Discussion Papers and Conference proceedings"
    learning_module_start_pattern = r"Learning Module:"


    # End markers for each block (these define where a block *must* end if found)
    # Order matters for these end_patterns: the earlier in the list, the higher its priority as a boundary.
    section_end_markers = [
        r"\n\nLanguage skills:", # Specific next section title with double newline
        r"\n\nExpert’s contact information", # Specific next section title with double newline
        r"\n\nCertification:", # Major section title
        r"\n\nCountries of Work Experience:", # Major section title
        r"\n\nEmployment record and relevant experiences to the assignment :", # Major section title
        r"\n\nEmployment record:", # Major section title
        r"\n\nRelevant Constancy and Research Experience:", # Major section title
        r"\Z" # End of document
    ]

    # --- Helper to extract content and the text remaining after the extraction ---
    def extract_block_and_remainder(text_to_search, start_re_pattern, next_section_start_patterns, fallback_end_patterns):
        start_match = re.search(start_re_pattern, text_to_search, re.DOTALL | re.IGNORECASE)
        if not start_match:
            return None, text_to_search # Block not found, return original text

        block_start_index = start_match.end()
        content_after_start = text_to_search[block_start_index:]
       
        # Find the earliest end point: either a specific next section start, or a general end marker
        earliest_end_index = len(content_after_start) # Default to end of available text

        # Prioritize specific next section start patterns
        for next_start_p in next_section_start_patterns:
            temp_match = re.search(next_start_p, content_after_start, re.DOTALL | re.IGNORECASE)
            if temp_match:
                earliest_end_index = min(earliest_end_index, temp_match.start())

        # Then consider general fallback end patterns
        for end_p in fallback_end_patterns:
            
            temp_match = re.search(end_p, content_after_start, re.DOTALL | re.IGNORECASE)
            if temp_match:
                earliest_end_index = min(earliest_end_index, temp_match.start())
        
        extracted_content = content_after_start[:earliest_end_index].strip()
        remaining_text = content_after_start[earliest_end_index:].strip() # Text after the extracted block
       
        return extracted_content, remaining_text

    current_text_to_parse = full_text # Start with the whole document text

    # 1. Extract Journals
    journals_content, current_text_to_parse = extract_block_and_remainder(
        current_text_to_parse,
        journals_start_pattern,
        [book_chapters_start_pattern, learning_module_start_pattern], # Specific next sections
        section_end_markers # General end markers
    )
    if journals_content:
        publications_data["publications"].append('journals')
        # Split by newlines and filter out empty lines. Assume each line is an entry.
        entries = [line.strip() for line in journals_content.split('\n') if line.strip()]
        publications_data["journals"] = entries
        
    
    # 2. Extract Book Chapters (search in the remaining text)
    book_chapters_content, current_text_to_parse = extract_block_and_remainder(
        current_text_to_parse,
        book_chapters_start_pattern,
        [learning_module_start_pattern], # Specific next section
        section_end_markers # General end markers
    )
    if book_chapters_content:
        publications_data['publications'].append('Book chapters, Discussion Papers and Conference proceedings')
        # Split by newlines and filter out empty lines.
        # No need to filter out internal headers if the block extraction is precise.
        entries = [line.strip() for line in book_chapters_content.split('\n') if line.strip()]
        publications_data["books"] = entries
        # publications_data["publications"].extend(entries)

    # 3. Extract Learning Module (search in the remaining text)
    learning_module_content, current_text_to_parse = extract_block_and_remainder(
        current_text_to_parse,
        learning_module_start_pattern,
        [r"Language skills:"], # No specific next publication section after learning module
        section_end_markers # General end markers
    )
    if learning_module_content:
        publications_data['publications'].append('learning module')
        entries = [line.strip() for line in learning_module_content.split('\n') if line.strip()]
        publications_data["learning_module"] = entries
        # publications_data["publications"].extend(entries)

    # Ensure uniqueness in the combined publications list while preserving order
    publications_data["publications"] = list(dict.fromkeys(publications_data["publications"]))

    return publications_data


def parse_full_cv(docx_path):
    """
    Parses a CV document (.docx) to extract structured information into a dictionary
    conforming to the provided Django model structure.
    """
    doc = Document(docx_path)
    # Read paragraphs and join them, ensuring stripped content and clean newlines
    full_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    # print(full_text)
    # Replace multiple newlines with single double newline for consistent sectioning
    full_text = re.sub(r'\n\s*\n+', '\n\n', full_text).strip()

    parsed = {
        "expert": {
            "first_name": None,
            "last_name": None,
            "email": None,
            "cv_language": "English",
            "country": None, # Will be filled from personal_detail
            "expertise_area": None,
            "language_skills": [],
            "publications": [], # Will be a list of strings
            "journals": [],     # Will be a list of strings
            "learning_module": [], # Will be a list of strings
            "books": [],        # Will be a list of strings
            "countries_of_work_experience": [] # New field to capture
        },
        "personal_detail": {
            "date_of_birth": None,
            "gender": None,
            "country": None,
            "phone_number": None,
            "email": None,
            "cv_language": "English",
            "current_position": None,
            "name_suffix": None
        },
        "education": [],
        "work_experience": [],
        "certifications": [],
        "research_experience": [],
    }

    # --- Extract Personal and Contact Information (prioritizing tables) ---
    personal_info = extract_personal_info(full_text, doc)
    
    # Update expert and personal_detail dictionaries using the extracted info
    parsed["expert"].update({
        "first_name": personal_info["first_name"],
        "last_name": personal_info["last_name"],
        "email": personal_info["email"],
        "phone_number": personal_info["phone_number"],
        "country": personal_info["country"] # Expert country derived from personal_info
    })
    # print(parsed['expert'])
    parsed["personal_detail"].update({
        "date_of_birth":personal_info["date_of_birth"],
        "country":personal_info["country"],
        "phone_number":personal_info["phone_number"],
        "email":personal_info["email"],
        "current_position":personal_info["current_position"],
        "name_suffix":personal_info["name_suffix"],

    })

   # Step 1: Match the section with the countries
    countries_work_exp_match = re.search(
        r"Countries of Work Experience:?\s*([A-Za-z,\s;-]+)",  # Keep it open-ended
        full_text, re.IGNORECASE | re.DOTALL
    )

    if countries_work_exp_match:
        countries_str = countries_work_exp_match.group(1).strip()

        # Step 2: Stop at certain keywords if they exist
        stop_words = ["employment", "education", "experience", "responsibilities"]
        stop_pattern = re.compile(rf"\b({'|'.join(stop_words)})\b", re.IGNORECASE)

        stop_match = stop_pattern.search(countries_str)
        if stop_match:
            countries_str = countries_str[:stop_match.start()].strip()

        # Step 3: Now safely split the trimmed string into a list
        countries_list = [
        c.strip() for c in re.split(r';\s*|,\s*|\s+and\s+', countries_str) if c.strip()
    ]

        parsed["expert"]["countries_of_work_experience"] = countries_list

    # --- Parse all other Tables ---
    # Iterate through tables to find sections by headers.
    # Note: We've already processed the first table (personal details) implicitly
    # in extract_personal_info, but iterating through all tables here is fine
    # as long as the header checks are distinct.
    for table_idx, table in enumerate(doc.tables):
        # Assuming tables that have headers relevant to these sections are not the initial personal info table
        # Check if the table has at least one row for headers
        if len(table.rows) == 0:
            continue
            
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        
        # Education Section (Title: "Education :", Headers: "Institution Name", "Degree Level", "Field of Study", "Year of Graduation")
        if "institution name" in headers and "degree level" in headers and "field of study" in headers:
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 4:
                  
                    parsed["education"].append({
                        "institution_name": cells[0],
                        "education_level": normalize_level(cells[1]),
                        "field_of_study": cells[2],
                        "year_of_grad": safe_date_parse(cells[3])
                    })
        
        # Professional Competency Certification Section (Title: "professional Competency Certification:", Headers: "Place/Organization", "Start Date", "End Date", "Field of training")
        elif any("place" in h for h in headers) and any("field of training" in h for h in headers):
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                
                place_org = cells[0].strip() if len(cells) > 0 else ""
                start_date = None
                end_date = None
                field_of_training = ""
                
                # Dynamic date parsing based on typical formats in this table
                # Handles "Month Day,", "Year", "Month Day,", "Year", "Field" (6 cells)
                if len(cells) >= 6 and re.match(r"^\d{4}$", cells[2]) and re.match(r"^\d{4}$", cells[4]):
                    start_date = safe_date_parse(f"{cells[1]} {cells[2]}")
                    end_date = safe_date_parse(f"{cells[3]} {cells[4]}")
                    field_of_training = cells[5].strip()
                # Handles "Start Date", "End Date", "Field" (4 cells)
                elif len(cells) >= 4:
                    start_date = safe_date_parse(cells[1])
                    end_date = safe_date_parse(cells[2])
                    field_of_training = cells[3].strip()
                
                if place_org and (start_date or end_date or field_of_training):
                    parsed["certifications"].append({
                        "place": place_org,
                        "start_date": start_date,
                        "end_date": end_date,
                        "field_of_training": field_of_training,
                        "typee": "certification", # Matches WorkExperience model's 'typee'
                        "description": field_of_training
                    })

        # Employment Record (Work Experience) Section (Title: "Employment record:", Headers: "Position Title", "Organization", "Start Date", "End Date", "country", "Responsibility")
    
        elif "position title" in headers and "organization" in headers and "responsibility" in headers:
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 6:
                    start_date_str = cells[2].strip()
                    end_date_str = cells[3].strip()

                    parsed["work_experience"].append({
                        "position_title": cells[0],
                        "organization_name": cells[1],
                        "start_date": safe_date_parse(start_date_str),
                        "end_date": safe_date_parse(end_date_str),
                        "country": cells[4].strip() if len(cells) > 4 and cells[4].strip() else None,
                        "responsibilities": cells[5].strip() if len(cells) > 5 else None,
                        "typee": "work_experience",
                        "description": None
                    })
        
        # Relevant Consultancy and Research Experience Section (Title: "Relevant Constancy and Research Experience:", Headers: "Period", "Employing organization and your title/position. Contact information for references", "Country", "Summary of activities performed relevant to the Assignment")
        elif "period" in headers and "summary of activities performed relevant to the assignment" in headers:
            current_category_buffer = [] 
            current_category = None # Initialize here, it will persist across project rows

            for i, row in enumerate(table.rows[1:]):
                cells = [c.text.strip() for c in row.cells]
                
                temp_period_cell_content = cells[0].strip() if cells else ""
                temp_org_contact_cell_content = cells[1].strip() if len(cells) > 1 else ""
                temp_country_cell_content = cells[2].strip() if len(cells) > 2 else ""
                temp_summary_activities_cell_content = cells[3].strip() if len(cells) > 3 else ""

                is_category_row = False
                
                # Your existing heuristic for `is_category_row` detection
                period_is_date_like = False
                if re.match(r"(\w+\s*\d{4}|\d{4}|present)\s*(?:-|to|continues up to)\s*(\w+\s*\d{4}|\d{4}|present)", temp_period_cell_content, re.IGNORECASE) or \
                   re.match(r"^\d{4}$", temp_period_cell_content) or \
                   re.match(r"^(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}$", temp_period_cell_content, re.IGNORECASE):
                    period_is_date_like = True
                    
                has_position_info = re.search(r"Position\s*:", temp_org_contact_cell_content, re.IGNORECASE)
                has_client_info = re.search(r"Client\s*:", temp_org_contact_cell_content, re.IGNORECASE)
                has_project_info = re.search(r"Project\s*:", temp_org_contact_cell_content, re.IGNORECASE)
                has_contact_info = re.search(r"(?:Email|Tel|Contact person)\s*:", temp_org_contact_cell_content, re.IGNORECASE)

                if temp_period_cell_content and \
                   (not period_is_date_like or not (has_position_info or has_client_info or has_project_info or has_contact_info)) and \
                   (not temp_country_cell_content or temp_country_cell_content.lower().strip() == temp_period_cell_content.lower().strip() or len(temp_country_cell_content.strip()) < 5) and \
                   (not temp_summary_activities_cell_content or temp_summary_activities_cell_content.lower().strip() == temp_period_cell_content.lower().strip() or len(temp_summary_activities_cell_content.strip()) < 5):
                    if len(cells) == 1 or all(c.lower().strip() == temp_period_cell_content.lower().strip() or not c.strip() for c in cells[1:min(len(cells),4)]):
                        is_category_row = True


                if is_category_row:
                    current_category_buffer.append(temp_period_cell_content) 
                    continue # Skip this row, as it's part of a multi-line header
                else: # This is a project row
                    # If there's accumulated content in the buffer, it means a new category just ended
                    if current_category_buffer: 
                        current_category = " ".join(current_category_buffer).strip()
                        current_category_buffer = [] # Clear buffer after using it for the current category
                    
                    # Now, use the (potentially new, or existing from a previous project) current_category for this project row
                    if len(cells) >= 4:
                        period_str = temp_period_cell_content
                        org_contact_str = temp_org_contact_cell_content
                        country_str = temp_country_cell_content
                        summary_activities = temp_summary_activities_cell_content

                        start_date, end_date = parse_date_range(period_str.strip())



                        client_match = re.search(r"Client\s*:\s*(.+?)(?=\s*(?:Position|Contact person|Email|Tel|$))", org_contact_str, re.IGNORECASE | re.DOTALL)
                        position_match = re.search(r"Position\s*:\s*(.+?)(?=\s*(?:Client|Contact person|Email|Tel|$))", org_contact_str, re.IGNORECASE | re.DOTALL)
                        contact_person_match = re.search(r"Contact person\s*:\s*(.+?)(?=\s*(?:Client|Position|Email|Tel|$))", org_contact_str, re.IGNORECASE | re.DOTALL)
                        email_in_cell_match = re.search(r"Email\s*:\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]+)", org_contact_str, re.IGNORECASE)
                        phone_in_cell_match = re.search(r"Tel\s*:\s*(\+?\d[\d\s-]{8,}\d)", org_contact_str, re.IGNORECASE)
                        
                        project_name = None
                        description_content = None
                        project_match_in_summary = re.search(r"Project\s*:\s*(.+?)(?:\nActivity performed:|\Z)", summary_activities, re.IGNORECASE | re.DOTALL)
                        activity_performed_match = re.search(r"Activity performed:\s*(.+)", summary_activities, re.IGNORECASE | re.DOTALL)

                        if project_match_in_summary:
                            project_name = project_match_in_summary.group(1).strip()
                        
                        if activity_performed_match:
                            description_content = activity_performed_match.group(1).strip()
                        
                        if not project_name and not description_content:
                            description_content = summary_activities.strip()
                        elif not description_content and project_name:
                            description_content = summary_activities.replace(f"Project: {project_name}", "").strip()

                        parsed["research_experience"].append({
                            "position": position_match.group(1).strip() if position_match else None,
                            "start_date": start_date, "end_date": end_date, "country": country_str,
                            "email": email_in_cell_match.group(1).strip() if email_in_cell_match else None,
                            "client": client_match.group(1).strip() if client_match else None,
                            "contact_person": contact_person_match.group(1).strip() if contact_person_match else None,
                            "phone_number": re.sub(r"\s+|-", "", phone_in_cell_match.group(1)) if phone_in_cell_match else None,
                            "description": description_content,
                            "project_name": project_name,
                            "category": current_category # This will now hold the last valid category
                        })
    
        
        # Language Skills Section (Headers: "Language", "Reading", "Speaking", "Writing")
        elif "language" in headers and "reading" in headers and "speaking" in headers and "writing" in headers:
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 4:
                    try:
                        parsed["expert"]["language_skills"].append({
                            "language": cells[0],
                            "reading": int(cells[1]),
                            "speaking": int(cells[2]),
                            "writing": int(cells[3])
                        })
                    except ValueError:
                        pass # Skip if scores are not valid integers (e.g., non-numeric data)

    # --- Publications, Books, Journals, Learning Modules (Non-Table Sections) ---

    publications_data = parse_publications(full_text)
    parsed["expert"]["publications"] = publications_data["publications"] # This is now a list
    parsed["expert"]["journals"] = publications_data["journals"]
    parsed["expert"]["books"] = publications_data["books"]
    parsed["expert"]["learning_module"] = publications_data["learning_module"]

    # Convert lists to '*' delimited strings for Django TextField
    # (This step assumes your Django model's get_books/set_books methods handle this format)
    parsed["expert"]["publications"] = '*'.join(parsed["expert"]["publications"])
    parsed["expert"]["journals"] = '*'.join(parsed["expert"]["journals"])
    parsed["expert"]["books"] = '*'.join(parsed["expert"]["books"])
    parsed["expert"]["learning_module"] = '*'.join(parsed["expert"]["learning_module"])

    # --- Expertise Area (derived from education fields) ---
    expertise_areas = [edu['field_of_study'] for edu in parsed["education"] if edu['field_of_study']]
    if expertise_areas:
        parsed["expert"]["expertise_area"] = ", ".join(list(set(expertise_areas))) # Use set to avoid duplicates

    return parsed