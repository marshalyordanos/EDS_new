import PyPDF2
import docx
import os

def extract_text(file, filename=None):
   """
   Extract text from PDF, DOCX, TXT using PyPDF2 and python-docx.
   `file` can be a file path (str) or a file-like object.
   `filename` is used to determine the file type if file-like object is passed.
   """
   try:
       # Determine extension
       if isinstance(file, str):
           file_ext = file.lower()
       elif filename:
           file_ext = filename.lower()
       else:
           print("No filename provided for file-like object")
           return ""
      
       if file_ext.endswith(".pdf"):
           if isinstance(file, str):
               with open(file, "rb") as f:
                   reader = PyPDF2.PdfReader(f)
                   return "\n".join(page.extract_text() or "" for page in reader.pages)
           else:
               reader = PyPDF2.PdfReader(file)
               return "\n".join(page.extract_text() or "" for page in reader.pages)

       elif file_ext.endswith(".docx"):
           if isinstance(file, str):
               doc = docx.Document(file)
           else:
               # python-docx requires a path or a BytesIO object
               from io import BytesIO
               doc = docx.Document(BytesIO(file.read()))
               file.seek(0)  # reset file pointer
           return "\n".join(para.text for para in doc.paragraphs)

       elif file_ext.endswith(".txt"):
           if isinstance(file, str):
               with open(file, "r", encoding="utf-8", errors="ignore") as f:
                   return f.read()
           else:
               return file.read().decode("utf-8", errors="ignore")

   except Exception as e:
       print(f"Error extracting text from {filename or file}: {e}")

   print("not successfully read.")
   return ""

